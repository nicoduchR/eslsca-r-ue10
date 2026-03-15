"""
Génération du rapport Word - Analyse Voitures d'occasion
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles globaux ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

plots = "resultats_analyse/plots"

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_img(doc, path, width=Inches(5.5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# =============================================================================
# PAGE DE TITRE
# =============================================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Analyse Statistique\nVoitures d'occasion")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
run.bold = True

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Dataset : data - voitures.csv\n98 observations - 7 variables")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("ESLSCA - UE10")
run.font.size = Pt(12)

doc.add_page_break()

# =============================================================================
# SOMMAIRE
# =============================================================================
doc.add_heading("Sommaire", level=1)
sommaire = [
    "1. Statistiques descriptives univariées",
    "2. Statistiques descriptives bivariées",
    "3. Tests statistiques",
    "    Test 1 : Corrélation de Pearson - Prix × Kilométrage",
    "    Test 2 : Régression multiple - Prix ~ Kilométrage + Année",
    "4. Le prix dépend du kilométrage mais pas de l'année : explication",
    "Annexe : Code R complet",
]
for item in sommaire:
    p = doc.add_paragraph(item)
    if item.startswith("    "):
        p.paragraph_format.left_indent = Cm(1.5)

doc.add_page_break()

# =============================================================================
# PARTIE 1 : UNIVARIÉES
# =============================================================================
doc.add_heading("1. Statistiques descriptives univariées", level=1)

# --- 1.1 Prix ---
doc.add_heading("1.1 Prix", level=2)

add_table(doc,
    ["Statistique", "Valeur"],
    [
        ["Moyenne", "15 165 EUR"],
        ["Médiane", "9 750 EUR"],
        ["Écart-type", "13 835 EUR"],
        ["CV", "91.2%"],
        ["Min - Max", "1 900 - 80 000 EUR"],
        ["Q1 - Q3", "5 600 - 20 000 EUR"],
    ])

add_img(doc, f"{plots}/hist_boxplot_prix.png")

p = doc.add_paragraph()
run = p.add_run("Interprétation : ")
run.bold = True
p.add_run(
    "La distribution du prix est fortement asymétrique à droite : la moyenne (15 165 EUR) "
    "est bien supérieure à la médiane (9 750 EUR), ce qui signifie que quelques véhicules "
    "très chers tirent la moyenne vers le haut. Le coefficient de variation de 91% traduit "
    "une dispersion extrême - le dataset mélange des véhicules à bas prix (vieilles citadines "
    "diesel autour de 2 000 EUR) et des véhicules haut de gamme (BMW Z1 à 80 000 EUR, "
    "Toyota Supra à 56 000 EUR). La médiane est plus représentative du véhicule « type »."
)

# --- 1.2 Kilométrage ---
doc.add_heading("1.2 Kilométrage", level=2)

add_table(doc,
    ["Statistique", "Valeur"],
    [
        ["Moyenne", "132 764 km"],
        ["Médiane", "108 000 km"],
        ["Écart-type", "100 881 km"],
        ["CV", "76%"],
        ["Min - Max", "100 - 498 000 km"],
        ["Q1 - Q3", "60 100 - 198 000 km"],
    ])

add_img(doc, f"{plots}/hist_boxplot_kilometrage.png")

p = doc.add_paragraph()
run = p.add_run("Interprétation : ")
run.bold = True
p.add_run(
    "Distribution également asymétrique à droite avec une forte dispersion (CV = 76%). "
    "On observe des extrêmes : un véhicule quasi-neuf (100 km) et des véhicules à très "
    "fort kilométrage (498 000 km). La moitié des véhicules se situe entre 60 000 et 198 000 km."
)

# --- 1.3 Énergie ---
doc.add_heading("1.3 Type d'énergie", level=2)

add_table(doc,
    ["Énergie", "Effectif", "%"],
    [
        ["Diesel", "38", "38.8%"],
        ["Essence", "28", "28.6%"],
        ["Hybride", "27", "27.6%"],
        ["Électrique", "5", "5.1%"],
    ])

add_img(doc, f"{plots}/barplot_energie.png")

p = doc.add_paragraph()
run = p.add_run("Interprétation : ")
run.bold = True
p.add_run(
    "Le diesel domine encore (39%), mais les motorisations hybrides représentent déjà 28% "
    "du parc, signe de la transition énergétique en cours. L'électrique reste marginal (5 véhicules)."
)

doc.add_page_break()

# =============================================================================
# PARTIE 2 : BIVARIÉES
# =============================================================================
doc.add_heading("2. Statistiques descriptives bivariées", level=1)

# --- 2.1 Prix × Km ---
doc.add_heading("2.1 Prix en fonction du kilométrage", level=2)

add_table(doc,
    ["Indicateur", "Valeur"],
    [
        ["Pearson r", "-0.601"],
        ["R²", "0.361 (36.1%)"],
        ["Équation", "prix = 26 108 − 0.0824 × km"],
    ])

add_img(doc, f"{plots}/scatter_prix_km.png")

p = doc.add_paragraph()
run = p.add_run("Interprétation : ")
run.bold = True
p.add_run(
    "Relation négative nette : plus le kilométrage augmente, plus le prix baisse. "
    "Le R² de 0.36 signifie que le kilométrage seul explique 36% de la variance du prix. "
    "En moyenne, chaque tranche de 10 000 km supplémentaires fait baisser le prix "
    "d'environ 824 EUR."
)

# --- 2.2 Prix × Année ---
doc.add_heading("2.2 Prix en fonction de l'année", level=2)

add_table(doc,
    ["Indicateur", "Valeur"],
    [
        ["Pearson r", "0.324"],
        ["R²", "0.105 (10.5%)"],
        ["Équation", "prix = −1 339 255 + 671.9 × année"],
    ])

add_img(doc, f"{plots}/scatter_prix_annee.png")

p = doc.add_paragraph()
run = p.add_run("Interprétation : ")
run.bold = True
p.add_run(
    "La corrélation est faible (r = 0.32) et le R² de seulement 10.5% montre que "
    "l'année seule explique peu le prix. Le nuage de points est très dispersé : on observe "
    "des véhicules anciens à prix élevés (BMW Z1 de 1990 à 80 000 EUR, Toyota Supra de "
    "1992 à 23 000 EUR) qui contredisent la tendance. L'année est un mauvais prédicteur "
    "du prix comparée au kilométrage."
)

# --- 2.3 Prix × Énergie ---
doc.add_heading("2.3 Prix selon le type d'énergie", level=2)

add_table(doc,
    ["Énergie", "n", "Moyenne", "Médiane", "Écart-type"],
    [
        ["Diesel", "38", "9 289", "5 600", "8 467"],
        ["Électrique", "5", "23 260", "28 000", "11 090"],
        ["Essence", "28", "16 764", "7 400", "19 567"],
        ["Hybride", "27", "20 278", "19 000", "10 130"],
    ])

add_img(doc, f"{plots}/boxplot_prix_energie.png")

p = doc.add_paragraph()
run = p.add_run("Interprétation : ")
run.bold = True
p.add_run(
    "Les véhicules hybrides et électriques sont en moyenne plus chers que les diesels. "
    "Les diesels ont la médiane la plus basse (5 600 EUR), reflétant leur ancienneté et "
    "la baisse de demande. L'essence présente une très forte dispersion à cause de véhicules "
    "de collection/sportifs."
)

doc.add_page_break()

# =============================================================================
# PARTIE 3 : TESTS STATISTIQUES
# =============================================================================
doc.add_heading("3. Tests statistiques", level=1)

# --- Test 1 ---
doc.add_heading("Test 1 : Corrélation de Pearson - Prix × Kilométrage", level=2)

p = doc.add_paragraph()
run = p.add_run("Hypothèses :")
run.bold = True

doc.add_paragraph("H0 : Il n'y a pas de corrélation linéaire entre le prix et le kilométrage (r = 0)", style='List Bullet')
doc.add_paragraph("H1 : Il existe une corrélation linéaire (r ≠ 0)", style='List Bullet')
doc.add_paragraph("Alpha = 0.05", style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Conditions de validité :")
run.bold = True

add_table(doc,
    ["Variable", "Shapiro-Wilk W", "p-value", "Normalité"],
    [
        ["Prix", "0.8022", "< 0.001", "Rejetée"],
        ["Kilométrage", "0.8762", "< 0.001", "Rejetée"],
    ])

doc.add_paragraph(
    "La normalité est rejetée pour les deux variables. Le test de Pearson est néanmoins "
    "appliqué car il reste robuste pour les grands échantillons (n = 98). "
    "Le test de Spearman est utilisé en complément."
)

add_img(doc, f"{plots}/qqplot_prix.png", width=Inches(4))
add_img(doc, f"{plots}/qqplot_kilometrage.png", width=Inches(4))

p = doc.add_paragraph()
run = p.add_run("Résultats :")
run.bold = True

add_table(doc,
    ["Méthode", "Coefficient", "Statistique", "p-value"],
    [
        ["Pearson", "r = -0.601", "t = -7.367, df = 96", "< 0.001"],
        ["Spearman", "rho = -0.826", "-", "< 0.001"],
    ])

p = doc.add_paragraph()
run = p.add_run("Conclusion : ")
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
p.add_run(
    "p < 0.05 → on rejette H0. Il existe une corrélation négative significative entre "
    "le prix et le kilométrage (r = -0.60, rho = -0.83, p < 0.001). Plus un véhicule "
    "a de kilomètres, plus son prix baisse. Le Spearman (rho = -0.83) plus élevé en "
    "valeur absolue que le Pearson (r = -0.60) indique une relation monotone forte "
    "mais non strictement linéaire."
)

doc.add_page_break()

# --- Test 2 ---
doc.add_heading("Test 2 : Régression multiple - Prix ~ Kilométrage + Année", level=2)

p = doc.add_paragraph()
run = p.add_run("Objectif : ")
run.bold = True
p.add_run(
    "Vérifier si l'année apporte une information supplémentaire sur le prix, "
    "au-delà de ce que le kilométrage explique déjà."
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Hypothèses :")
run.bold = True

doc.add_paragraph("H0 : Le coefficient de l'année est nul (β_année = 0) quand le kilométrage est déjà dans le modèle", style='List Bullet')
doc.add_paragraph("H1 : Le coefficient de l'année est significativement différent de 0", style='List Bullet')
doc.add_paragraph("Alpha = 0.05", style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Prérequis - colinéarité kilométrage/année :")
run.bold = True

doc.add_paragraph(
    "Les véhicules plus anciens ont plus de kilomètres : la corrélation entre kilométrage "
    "et année est r = -0.52 (p < 0.001). Les deux variables portent une information "
    "en partie redondante."
)

add_img(doc, f"{plots}/scatter_km_annee.png")

p = doc.add_paragraph()
run = p.add_run("Résultats de la régression multiple :")
run.bold = True

add_table(doc,
    ["Variable", "Coefficient", "t", "p-value", "Significatif ?"],
    [
        ["Kilométrage", "-0.0813", "-6.173", "< 0.001", "Oui"],
        ["Année", "+32.1", "0.162", "0.872", "Non"],
    ])

add_table(doc,
    ["Modèle", "R²"],
    [
        ["prix ~ kilométrage (seul)", "0.3612"],
        ["prix ~ kilométrage + année", "0.3613"],
        ["Gain en ajoutant l'année", "0.0001 (quasi nul)"],
    ])

doc.add_paragraph(
    "Comparaison des modèles (test F) : F = 0.026, p = 0.872 → le modèle avec année "
    "n'est pas significativement meilleur."
)

p = doc.add_paragraph()
run = p.add_run("Conclusion : ")
run.bold = True
run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
p.add_run(
    "p = 0.872 → on ne rejette pas H0. L'année n'apporte aucune information supplémentaire "
    "significative sur le prix une fois le kilométrage pris en compte. Le R² n'augmente "
    "quasiment pas (de 0.3612 à 0.3613). Le kilométrage reste le seul prédicteur significatif."
)

doc.add_page_break()

# =============================================================================
# PARTIE 4 : EXPLICATION ENRICHIE
# =============================================================================
doc.add_heading("4. Le prix dépend du kilométrage mais pas de l'année", level=1)

# --- Le constat chiffré ---
doc.add_heading("Le constat chiffré", level=2)

add_table(doc,
    ["Modèle", "Kilométrage", "Année", "R²"],
    [
        ["prix ~ km", "p < 0.001", "-", "0.3612"],
        ["prix ~ km + année", "p < 0.001", "p = 0.872", "0.3613"],
    ])

doc.add_paragraph(
    "Quand on connaît le kilométrage d'un véhicule, ajouter l'année au modèle ne change "
    "strictement rien : le R² passe de 0.3612 à 0.3613 - un gain de 0.01%, autant dire zéro."
)

# --- Raison 1 ---
doc.add_heading("Raison 1 : Le kilométrage mesure l'usure réelle, l'année ne mesure rien de concret", level=2)

doc.add_paragraph(
    "Ce qui fait perdre de la valeur à un véhicule, c'est son usure mécanique : moteur, "
    "embrayage, transmission, suspension, freins. Cette usure est directement proportionnelle "
    "aux kilomètres parcourus. Un véhicule qui a roulé 200 000 km a subi une dégradation "
    "physique mesurable, quel que soit son âge."
)
doc.add_paragraph(
    "L'année, en revanche, ne cause pas directement d'usure. Une voiture stockée dans un "
    "garage pendant 10 ans sans rouler ne perd pas de valeur mécanique."
)

# --- Raison 2 ---
doc.add_heading("Raison 2 : L'année est un simple intermédiaire (variable confondante)", level=2)

doc.add_paragraph(
    "Si l'année semble corrélée au prix en analyse simple (r = 0.32, p = 0.001), c'est "
    "uniquement parce que :"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Année ancienne → plus de temps pour rouler → plus de km → prix plus bas")
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph(
    "L'année n'agit pas sur le prix : elle agit sur le kilométrage, qui lui agit sur le prix. "
    "C'est une corrélation indirecte, pas une causalité. La régression multiple le démontre : "
    "une fois le kilométrage connu, l'année n'apporte plus aucune information (p = 0.872)."
)

# --- Raison 3 ---
doc.add_heading("Raison 3 : La démonstration par les chiffres", level=2)

doc.add_paragraph(
    "Le kilométrage et l'année sont corrélés (r = -0.52, p < 0.001). Cette corrélation "
    "crée une illusion : quand on teste l'année seule face au prix, on trouve une "
    "corrélation positive (r = 0.32, p = 0.001) qui semble significative. Mais cette "
    "corrélation est artificielle - elle n'existe que parce que les véhicules récents "
    "ont moins de kilomètres."
)
doc.add_paragraph(
    "La régression multiple le prouve : une fois qu'on connaît le kilométrage d'un "
    "véhicule, savoir son année de mise en circulation n'aide plus du tout à prédire "
    "son prix (p = 0.872). Toute l'information utile contenue dans l'année était déjà "
    "capturée par le kilométrage."
)

# --- Les cas concrets ---
doc.add_heading("Les cas concrets qui le confirment", level=2)

doc.add_paragraph(
    "Le dataset contient des exemples parlants :"
)

add_table(doc,
    ["Véhicule", "Année", "Km", "Prix", "Logique « année »", "Logique « km »"],
    [
        ["BMW Z1", "1990", "16 000", "80 000 EUR", "Devrait être quasi-nul", "Cohérent (peu roulé)"],
        ["Toyota Supra", "1992", "89 000", "23 000 EUR", "Devrait être très bas", "Cohérent (km modéré)"],
        ["Audi A2", "2013", "487 000", "5 400 EUR", "Devrait valoir plus", "Cohérent (km extrême)"],
    ])

doc.add_paragraph(
    "Ces véhicules contredisent la logique « plus ancien = moins cher », mais sont "
    "parfaitement cohérents avec la logique « plus de km = moins cher » (à l'exception "
    "des véhicules de collection dont la rareté prend le dessus)."
)

# --- L'analogie ---
doc.add_heading("L'analogie", level=2)

p = doc.add_paragraph()
p.add_run(
    "C'est comme demander : « Qu'est-ce qui détermine la fatigue d'un coureur : la "
    "distance parcourue ou le temps écoulé ? » La réponse est la distance. Le temps "
    "écoulé est corrélé à la distance (plus on court longtemps, plus on a parcouru), "
    "mais c'est bien la distance qui fatigue physiquement, pas le passage du temps."
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    "De la même façon, c'est le kilométrage (l'usure réelle) qui détermine le prix, "
    "pas l'année (le temps écoulé). Connaître la distance rend le temps inutile - "
    "exactement comme le kilométrage rend l'année inutile pour prédire le prix."
)
run.bold = True

doc.add_page_break()

# =============================================================================
# ANNEXE : CODE R
# =============================================================================
doc.add_heading("Annexe : Code R complet", level=1)

doc.add_paragraph(
    "Le script ci-dessous est reproductible. Pour l'exécuter :"
)
add_code(doc, "Rscript scripts/analyse_voitures.R")
doc.add_paragraph()

# Lire le fichier R
with open("scripts/analyse_voitures.R", "r") as f:
    r_code = f.read()

# Découper en blocs pour éviter un seul paragraphe immense
blocks = r_code.split("\n\n")
for block in blocks:
    if block.strip():
        add_code(doc, block.strip())

# =============================================================================
# SAUVEGARDE
# =============================================================================
output_path = "resultats_analyse/rapport_analyse_voitures.docx"
doc.save(output_path)
print(f"Document Word généré : {output_path}")
